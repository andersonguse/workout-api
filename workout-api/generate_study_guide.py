from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


out = Path(r"C:\Users\ander\Documents\workout-api\workout-api\Spring Boot Workout API Study Guide.docx")

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Aptos Display"
    styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

if "Code Block" not in styles:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)


def shade_paragraph(paragraph, fill="F2F2F2"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def title(text, subtitle=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(12)
    doc.add_paragraph()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text=""):
    doc.add_paragraph(text)


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(text):
    for line in text.strip("\n").split("\n"):
        para = doc.add_paragraph(line, style="Code Block")
        shade_paragraph(para)


def table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = head
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


title(
    "Spring Boot Workout API Study Guide",
    "Interview prep notes for your layered REST API project",
)

h1("1. Project Snapshot")
p(
    "This guide summarizes the Spring Boot workout API you built and gives you language to "
    "explain the architecture in an interview. The project is a Java, Maven, Spring Boot REST "
    "API backed by Dockerized PostgreSQL."
)
bullets(
    [
        "Java 25 with Maven",
        "Spring Boot 4.1.0",
        "Spring Web MVC for REST endpoints",
        "Spring Data JPA for database access",
        "PostgreSQL running in Docker",
        "Jakarta Validation for request validation",
        "Lombok for boilerplate getters, setters, and constructors",
        "Postman for endpoint testing",
    ]
)

h2("Current Feature Set")
table(
    ["Capability", "Endpoint", "Status Code Goal", "Purpose"],
    [
        ["Get all workouts", "GET /api/workouts", "200 OK", "Return all workouts as response DTOs"],
        ["Get one workout", "GET /api/workouts/{id}", "200 OK or 404 Not Found", "Return one workout by id"],
        ["Create workout", "POST /api/workouts", "201 Created or 400 Bad Request", "Validate request, save entity, return response DTO"],
        ["Update workout", "PUT /api/workouts/{id}", "200 OK or 404 Not Found", "Full update of an existing workout"],
        ["Delete workout", "DELETE /api/workouts/{id}", "204 No Content or 404 Not Found", "Remove an existing workout"],
    ],
)

h1("2. The Big Architecture Picture")
p(
    "Your API uses a layered architecture. Each layer has a focused responsibility, and "
    "dependencies flow downward from HTTP to business logic to persistence."
)
code(
    """
HTTP Request
    -> Controller / Presentation Layer
    -> Service / Business Layer
    -> Repository / Persistence Layer
    -> Database
"""
)
table(
    ["Layer", "Package", "Main Classes", "Responsibility"],
    [
        ["Presentation", "controller", "WorkoutController", "Accept HTTP requests, read path/body data, return HTTP responses"],
        ["Business / Service", "service", "WorkoutService", "Coordinate app behavior, enforce business flow, map DTOs/entities, call repository"],
        ["Persistence / Data Access", "repository, entity", "WorkoutRepository, Workout", "Represent database records and perform CRUD operations through JPA"],
        ["API Contract", "dto", "WorkoutRequest, WorkoutResponse", "Define what clients send and receive"],
        ["Error Handling", "exception", "WorkoutNotFoundException, GlobalExceptionHandler", "Convert application errors into clean HTTP responses"],
    ],
)
h2("Interview Summary")
code(
    """
I structured the app as a layered Spring Boot REST API. Controllers handle HTTP concerns,
services contain application logic and DTO mapping, repositories handle persistence through
Spring Data JPA, and entities map to database tables. Request DTOs define the API input
contract with validation, response DTOs define the output contract, and a global exception
handler converts domain exceptions into proper HTTP responses like 404 Not Found.
"""
)

h1("3. MVC vs REST API Layers")
p(
    "Classic MVC stands for Model, View, Controller. In a server-rendered web app, the view "
    "is usually HTML. In your REST API, the API returns JSON instead of HTML, so the practical "
    "structure is more often controller-service-repository."
)
table(
    ["MVC Term", "REST API Equivalent", "In Your Project"],
    [
        ["Model", "Domain/entity/DTO data structures", "Workout, WorkoutRequest, WorkoutResponse"],
        ["View", "JSON response body", "WorkoutResponse serialized to JSON"],
        ["Controller", "REST controller endpoint class", "WorkoutController"],
    ],
)

h1("4. Controller Layer")
p(
    "The controller is the presentation layer. It knows about HTTP routes, methods, path "
    "variables, request bodies, validation, and response status codes. It should not contain "
    "database logic."
)
h2("Important Controller Annotations")
table(
    ["Annotation", "Meaning", "Example"],
    [
        ["@RestController", "Marks a class as a REST controller. Return values are serialized as JSON.", "WorkoutController"],
        ["@RequestMapping", "Sets a base route for the controller.", '@RequestMapping("/api/workouts")'],
        ["@GetMapping", "Handles HTTP GET requests.", "GET /api/workouts"],
        ["@PostMapping", "Handles HTTP POST requests.", "POST /api/workouts"],
        ["@PutMapping", "Handles HTTP PUT requests.", "PUT /api/workouts/{id}"],
        ["@DeleteMapping", "Handles HTTP DELETE requests.", "DELETE /api/workouts/{id}"],
        ["@PathVariable", "Reads a value from the URL path.", "@PathVariable Long id"],
        ["@RequestBody", "Deserializes JSON body into a Java object.", "@RequestBody WorkoutRequest request"],
        ["@Valid", "Triggers Jakarta Bean Validation on the request body.", "@Valid @RequestBody WorkoutRequest request"],
    ],
)
h2("ResponseEntity")
p("ResponseEntity lets you explicitly control HTTP status codes and response bodies.")
code(
    """
ResponseEntity.ok(response)                         // 200 OK with body
ResponseEntity.status(HttpStatus.CREATED).body(x)   // 201 Created with body
ResponseEntity.noContent().build()                  // 204 No Content, no body
"""
)
code(
    """
I use ResponseEntity when I want explicit control over status codes. GET and PUT return
200 OK, POST returns 201 Created, and DELETE returns 204 No Content because there is no
response body.
"""
)

h1("5. Service Layer")
p(
    "The service layer is the business/application layer. In your app, it coordinates "
    "repository calls and maps between DTOs and entities. As the project grows, this is "
    "where rules belong: for example, endTime must be after startTime or totalWorkSeconds "
    "must match the duration."
)
h2("Constructor Injection")
code(
    """
private final WorkoutRepository workoutRepository;

public WorkoutService(WorkoutRepository workoutRepository) {
    this.workoutRepository = workoutRepository;
}
"""
)
p(
    "Constructor injection means the service cannot exist without its required dependency. "
    "The final keyword means the dependency must be initialized in the constructor and cannot "
    "be reassigned later."
)
h2("DTO Mapping in the Service")
code("WorkoutRequest -> Workout entity -> repository.save(workout) -> WorkoutResponse")
p(
    "Repositories save entities, not DTOs. WorkoutRepository extends JpaRepository<Workout, Long>, "
    "so save expects a Workout."
)
h2("Streams and map")
code(
    """
workoutRepository.findAll()
        .stream()
        .map(workout -> new WorkoutResponse(...))
        .toList();
"""
)
table(
    ["Step", "Type", "Meaning"],
    [
        ["findAll()", "List<Workout>", "Repository returns database entities"],
        ["stream()", "Stream<Workout>", "Creates a processing pipeline over the list"],
        ["map(...)", "Stream<WorkoutResponse>", "Converts each Workout into a WorkoutResponse"],
        ["toList()", "List<WorkoutResponse>", "Collects the mapped results into a list"],
    ],
)
p(
    "JavaScript arrays have map directly. Java List does not. Java uses stream().map().toList() "
    "because streams are a separate processing API."
)

h1("6. Repository Layer and Spring Data JPA")
p(
    "The repository is the persistence layer. It abstracts database access so you do not "
    "manually write SQL for common CRUD operations."
)
code(
    """
public interface WorkoutRepository extends JpaRepository<Workout, Long> {
}
"""
)
h2("Understanding JpaRepository<Workout, Long>")
table(
    ["Part", "Meaning"],
    [
        ["Workout", "The entity type managed by this repository"],
        ["Long", "The type of the entity primary key"],
    ],
)
p(
    "Because your entity has private Long id, the repository uses Long as the ID type. This "
    "gives you typed methods like findById(Long id), deleteById(Long id), and existsById(Long id)."
)
h2("Built-in Repository Methods")
table(
    ["Method", "Purpose"],
    [
        ["findAll()", "Return all rows as a List<Workout>"],
        ["findById(Long id)", "Return Optional<Workout> because the row may not exist"],
        ["save(Workout workout)", "Insert or update an entity"],
        ["deleteById(Long id)", "Delete by primary key"],
        ["existsById(Long id)", "Check whether a row exists"],
        ["count()", "Count rows"],
    ],
)
h2("Optional<Workout>")
p("findById returns Optional<Workout>, not Workout, because an ID may not exist.")
code(
    """
Workout workout = workoutRepository.findById(id)
        .orElseThrow(() -> new WorkoutNotFoundException(id));
"""
)

h1("7. Entity Layer")
p("The entity represents the database table. A Workout object maps to one row in the workouts table.")
code(
    """
Workout Java object  <-> workouts database table
One Workout instance <-> one database row
Workout fields       <-> table columns
"""
)
h2("Important Entity Annotations")
table(
    ["Annotation", "Meaning"],
    [
        ["@Entity", "Marks this class as a JPA entity managed by Hibernate"],
        ['@Table(name = "workouts")', "Maps the entity to the workouts table"],
        ["@Id", "Marks the primary key field"],
        ["@GeneratedValue(strategy = GenerationType.IDENTITY)", "Database generates the ID value"],
    ],
)
h2("Long vs long, Integer vs int")
p(
    "In JPA entities and DTOs, wrapper types like Long and Integer are often better than "
    "primitives like long and int because wrappers can be null."
)
table(
    ["Type", "Can Be Null?", "Default Value", "Why It Matters"],
    [
        ["long", "No", "0", "Cannot represent an unsaved/missing ID"],
        ["Long", "Yes", "null", "ID can be null before database generation"],
        ["int", "No", "0", "Cannot distinguish missing from intentionally zero"],
        ["Integer", "Yes", "null", "Can represent missing input or unset DB value"],
    ],
)
code(
    """
I prefer wrapper types like Long and Integer in entities and DTOs because they support null.
That lets an ID be null before persistence and lets the API distinguish missing input from
a real value like 0.
"""
)

h1("8. DTOs: Request and Response")
p(
    "DTO means Data Transfer Object. DTOs define what crosses the API boundary. They protect "
    "your persistence model from becoming your public API contract."
)
table(
    ["Class", "Direction", "Purpose", "Includes ID?"],
    [
        ["WorkoutRequest", "Client -> API", "Input payload for create/update with validation", "No"],
        ["WorkoutResponse", "API -> Client", "Output payload returned to clients", "Yes"],
        ["Workout", "Service/Repository -> Database", "JPA entity persisted by Hibernate", "Yes"],
    ],
)
h2("Save First, Then Build Response")
code(
    """
Workout savedWorkout = workoutRepository.save(workout);

return new WorkoutResponse(
        savedWorkout.getId(),
        savedWorkout.getStartTime(),
        savedWorkout.getEndTime(),
        savedWorkout.getSets(),
        savedWorkout.getTotalRestSeconds(),
        savedWorkout.getTotalWorkSeconds()
);
"""
)
p(
    "The response should represent what was actually persisted, not just what the client "
    "requested. If save fails, Spring/JPA throws an exception and the response is not built."
)

h1("9. Validation")
p(
    "Validation belongs mainly on request DTOs because they define the API input contract. "
    "The entity focuses on persistence mapping; the response DTO focuses on output shape."
)
h2("Common Jakarta Validation Annotations")
table(
    ["Annotation", "Meaning", "Example Use"],
    [
        ["@NotNull", "Value cannot be null", "startTime, endTime"],
        ["@NotBlank", "String cannot be null, empty, or whitespace", "name"],
        ["@Positive", "Number must be greater than 0", "sets, totalWorkSeconds"],
        ["@PositiveOrZero", "Number must be 0 or greater", "totalRestSeconds"],
        ["@Min(1)", "Number must be at least 1", "sets"],
        ["@Max(100)", "Number must be at most 100", "difficulty or rating"],
        ["@Size(min=3, max=100)", "String or collection length range", "name or notes"],
        ["@Email", "String must be email format", "email field"],
        ["@PastOrPresent", "Date/time must not be future", "completed workout date"],
        ["@Future", "Date/time must be future", "scheduled workout date"],
    ],
)
code(
    """
@PostMapping
public ResponseEntity<WorkoutResponse> createWorkout(
        @Valid @RequestBody WorkoutRequest request) {
    ...
}
"""
)
p("@Valid tells Spring to check WorkoutRequest before the service method runs. Invalid input returns 400 Bad Request.")

h1("10. Exception Handling")
p(
    "Your app uses a custom exception and global exception handler for not-found cases. "
    "This keeps controllers clean and avoids duplicating try/catch blocks."
)
h2("Custom Exception")
code(
    """
public class WorkoutNotFoundException extends RuntimeException {
    public WorkoutNotFoundException(Long id) {
        super("Workout not found with id: " + id);
    }
}
"""
)
h2("Global Exception Handler")
code(
    """
@RestControllerAdvice
public class GlobalExceptionHandler {

    @ExceptionHandler(WorkoutNotFoundException.class)
    public ResponseEntity<String> handleWorkoutNotFound(WorkoutNotFoundException exception) {
        return new ResponseEntity<>(exception.getMessage(), HttpStatus.NOT_FOUND);
    }
}
"""
)
table(
    ["Annotation", "Meaning"],
    [
        ["@RestControllerAdvice", "Global error handling for REST controllers"],
        ["@ExceptionHandler(WorkoutNotFoundException.class)", "Run this method when that exception is thrown"],
        ["ResponseEntity<String>", "Return a body and status code"],
        ["HttpStatus.NOT_FOUND", "HTTP 404"],
    ],
)
code(
    """
The service throws a domain-specific exception when a workout is missing. A global exception
handler catches that exception and converts it into a 404 response. That keeps controllers
focused on successful HTTP flow and centralizes error behavior.
"""
)

h1("11. PostgreSQL and Docker")
p(
    "Your database runs in Docker, while Spring Boot runs on your machine. Spring connects "
    "to Postgres using a JDBC URL."
)
code(
    """
Postman -> http://localhost:8080/api/workouts -> Spring Boot
Spring Boot -> jdbc:postgresql://localhost:5432/workout_db -> PostgreSQL
DBeaver -> localhost:5432 -> PostgreSQL
"""
)
h2("application.properties")
code(
    """
spring.datasource.url=jdbc:postgresql://localhost:5432/workout_db
spring.datasource.username=workout_user
spring.datasource.password=workout_password
spring.datasource.driver-class-name=org.postgresql.Driver

spring.jpa.hibernate.ddl-auto=update
spring.jpa.show-sql=true
spring.jpa.properties.hibernate.format_sql=true
"""
)
bullets(
    [
        "localhost:8080 is the Spring Boot API, not the database.",
        "localhost:5432 is PostgreSQL.",
        "DBeaver should connect using host localhost, port 5432, database workout_db.",
        "Postgres handles auto-increment IDs even when Spring Boot is not running.",
        "ddl-auto=update is convenient for learning, but production systems usually use Flyway or Liquibase.",
    ]
)

h1("12. Lombok")
p("Lombok reduces boilerplate Java code by generating methods during compilation.")
table(
    ["Annotation", "Meaning"],
    [
        ["@Getter", "Generates getter methods"],
        ["@Setter", "Generates setter methods"],
        ["@NoArgsConstructor", "Generates a no-argument constructor required by JPA"],
        ["@AllArgsConstructor", "Generates constructor with all fields"],
    ],
)
p("Avoid @Data on JPA entities because it generates equals, hashCode, and toString, which can cause problems when entities have IDs or relationships.")

h1("13. Testing Strategy")
p("You do not need every test written before the interview, but you should be able to explain how you would test each layer.")
table(
    ["Test Type", "Spring/JUnit Tool", "Purpose"],
    [
        ["Service unit tests", "JUnit 5 + Mockito", "Test business logic and DTO mapping without a real database"],
        ["Controller tests", "MockMvc / @WebMvcTest", "Test HTTP routes, status codes, request validation, JSON responses"],
        ["Repository tests", "@DataJpaTest", "Test JPA mappings and repository queries against a test database"],
        ["Integration tests", "@SpringBootTest", "Start more of the application and test real layer interaction"],
    ],
)
code(
    """
I would unit test the service layer with JUnit and Mockito by mocking WorkoutRepository.
I would test controller behavior with MockMvc to verify routes, status codes, and validation.
For persistence, I would use @DataJpaTest, and for full end-to-end flows I would use
@SpringBootTest.
"""
)

h1("14. Common Interview Questions and Answers")
qa = [
    ("What does @SpringBootApplication do?", "It marks the main class as the entry point for Spring Boot. It combines configuration, component scanning, and auto-configuration so Spring can discover controllers, services, repositories, and other beans."),
    ("What is dependency injection?", "Dependency injection means classes receive the objects they depend on instead of creating them manually. In this project, WorkoutService receives WorkoutRepository through its constructor, and WorkoutController receives WorkoutService."),
    ("Why use a service layer?", "The service layer keeps business/application logic out of controllers and persistence details out of the web layer. Controllers handle HTTP, services coordinate behavior, repositories handle database access."),
    ("Why use DTOs?", "DTOs separate the API contract from the database entity. Request DTOs define validated input, response DTOs define controlled output, and entities remain focused on persistence."),
    ("What is Spring Data JPA?", "Spring Data JPA provides repository abstractions over JPA/Hibernate. By extending JpaRepository, I get common CRUD methods without writing boilerplate SQL."),
    ("What is Hibernate?", "Hibernate is the JPA implementation that maps Java entities to database tables and translates repository operations into SQL."),
    ("What is Optional?", "Optional represents a value that may or may not exist. findById returns Optional because the requested row might be missing."),
    ("What is @RestControllerAdvice?", "It centralizes exception handling for REST controllers. It lets me convert exceptions into consistent HTTP responses."),
    ("Why use ResponseEntity?", "ResponseEntity lets me control HTTP status codes and response bodies explicitly, such as 201 Created for POST and 204 No Content for DELETE."),
    ("What is PUT vs PATCH?", "PUT is usually a full replacement/update of a resource. PATCH is usually a partial update where only provided fields change. This project currently uses PUT for full update."),
]
for question, answer in qa:
    h3(question)
    p(answer)

h1("15. Next Concepts To Learn")
nums(
    [
        "Extract mapping helper methods in WorkoutService: mapToResponse and applyRequestToWorkout.",
        "Add a WorkoutSet entity and model a one-to-many relationship from Workout to WorkoutSet.",
        "Learn @OneToMany, @ManyToOne, @JoinColumn, and cascade behavior carefully.",
        "Improve validation error responses in GlobalExceptionHandler.",
        "Add one service unit test with JUnit and Mockito.",
        "Learn spec driven development and write a short API spec for this project.",
    ]
)

h1("16. Spec Driven Development Connection")
p(
    "Since the hiring manager mentioned spec driven development, connect it to this API. "
    "Spec driven development means defining expected behavior before or while implementing. "
    "For an API, that can mean writing endpoint contracts, request/response examples, status "
    "codes, and acceptance criteria first."
)
h2("Example Spec for Create Workout")
code(
    """
Endpoint: POST /api/workouts
Request body:
{
  "startTime": "2026-06-15T07:30:00",
  "endTime": "2026-06-15T08:20:00",
  "sets": 14,
  "totalRestSeconds": 720,
  "totalWorkSeconds": 2280
}

Expected success:
201 Created
Response includes generated id and workout fields.

Expected validation failure:
400 Bad Request when required fields are missing or numeric values are invalid.
"""
)
code(
    """
For spec driven development, I would define the API behavior first: endpoint paths, methods,
request DTO shape, response DTO shape, status codes, validation rules, and error cases. Then
I would implement the controller, service, repository, and tests against that spec.
"""
)

h1("17. Quick Whiteboard Version")
code(
    """
POST /api/workouts
    JSON request body
        -> WorkoutRequest DTO (@Valid)
            -> WorkoutController
                -> WorkoutService
                    -> map request DTO to Workout entity
                    -> WorkoutRepository.save(entity)
                    -> map saved entity to WorkoutResponse DTO
                -> ResponseEntity 201 Created
"""
)

h1("18. Things To Clean Up After The Interview")
bullets(
    [
        "Use lowercase package names: com.anderson.workoutapi.",
        "Remove validation annotations from the entity if DTO validation is enough.",
        "Remove unused imports like LocalDateTime in the service and Workout in the controller.",
        "Extract duplicate DTO mapping into private helper methods or a mapper class.",
        "Add a real validation error response body.",
        "Use Flyway or Liquibase instead of ddl-auto=update for production-style schema management.",
        "Add tests for service, controller, and repository layers.",
    ]
)

p()
p(
    "End of guide. Study this as a story: HTTP comes in, controller delegates, service applies "
    "app flow and mapping, repository persists entities, exceptions and validation produce clean "
    "API behavior."
)

doc.save(out)
print(out)
